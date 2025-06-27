from PySide6.QtWebEngineCore import QWebEngineSettings
from docx import Document
from PySide6.QtWebEngineWidgets import QWebEngineView
from PySide6.QtCore import QUrl
import tempfile

class read_word():
    def __init__(self, docx_path):
        self.docx_path = docx_path

    # 将word文件转换成html 用于前面渲染
    def render_word_to_webview(self):
        doc = Document(self.docx_path)
        html = "<html><body>"
        # 文本段落
        for para in doc.paragraphs:
            html += f"<p>{para.text}</p>"

        # 表格
        for table in doc.tables:
            html += "<table border='1' style='border-collapse:collapse;'>"
            for row in table.rows:
                html += "<tr>"
                for cell in row.cells:
                    html += f"<td>{cell.text.strip()}</td>"
                html += "</tr>"
            html += "</table><br>"
        html += "</body></html>"

        webview = QWebEngineView()
        webview.settings().setAttribute(QWebEngineSettings.LocalContentCanAccessRemoteUrls, True)
        webview.settings().setAttribute(QWebEngineSettings.LocalContentCanAccessFileUrls, True)
        webview.setHtml(html)
        return webview

