import sys
import os
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from PySide6.QtWidgets import (
    QApplication, QWidget, QPushButton, QVBoxLayout,
    QFileDialog, QLabel, QMessageBox, QHBoxLayout, QFrame
)
from PySide6.QtGui import QFont, QColor, QPalette, QIcon
from PySide6.QtCore import Qt


# 设置中文字体
def set_chinese_font(run, font_name="微软雅黑", size=12):
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font_name)


def add_para(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    set_chinese_font(run)
    return para


def markdown_to_docx(md_text, output_file):
    html = markdown.markdown(md_text)
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for elem in soup.contents:
        if elem.name == 'h1':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'p':
            add_para(doc, elem.get_text())
        elif elem.name == 'ul':
            for li in elem.find_all('li'):
                add_para(doc, li.get_text(), style='ListBullet')
        elif elem.name == 'ol':
            for li in elem.find_all('li'):
                add_para(doc, li.get_text(), style='ListNumber')
        elif elem.name == 'blockquote':
            para = add_para(doc, elem.get_text())
            para.style = 'Intense Quote'
        elif elem.name == 'pre':
            para = add_para(doc, elem.get_text())
            para.style = 'Intense Quote'

    doc.save(output_file)


class MarkdownToWordApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Markdown 转 Word 工具")
        self.setWindowIcon(QIcon("keepnotesPlus.ico"))
        self.setFixedSize(500, 300)
        self.setStyleSheet("background-color: #f5f7fa; border-radius: 12px;")

        font = QFont("微软雅黑", 10)
        self.setFont(font)

        self.md_file = ""
        self.output_file = ""

        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # 文件标签
        self.label_md = QLabel("📝 未选择 Markdown 文件")
        self.label_out = QLabel("📁 未选择输出位置")
        self.label_md.setWordWrap(True)
        self.label_out.setWordWrap(True)

        # 按钮
        self.btn_select_md = QPushButton("选择 Markdown 文件")
        self.btn_select_md.clicked.connect(self.select_md_file)

        self.btn_select_output = QPushButton("选择输出位置")
        self.btn_select_output.clicked.connect(self.select_output_path)

        self.btn_start = QPushButton("🚀 开始转换")
        self.btn_start.clicked.connect(self.convert)

        for btn in [self.btn_select_md, self.btn_select_output, self.btn_start]:
            btn.setStyleSheet("""
                QPushButton {
                    background-color: #4a90e2;
                    color: white;
                    padding: 8px;
                    border-radius: 8px;
                }
                QPushButton:hover {
                    background-color: #357ABD;
                }
            """)

        # 美化容器（加边框）
        frame = QFrame()
        frame.setStyleSheet("""
            QFrame {
                background-color: white;
                border-radius: 10px;
                border: 1px solid #d0d0d0;
            }
        """)
        inner_layout = QVBoxLayout(frame)
        inner_layout.addWidget(self.label_md)
        inner_layout.addWidget(self.btn_select_md)
        inner_layout.addWidget(self.label_out)
        inner_layout.addWidget(self.btn_select_output)
        inner_layout.addWidget(self.btn_start)

        layout.addWidget(frame)
        self.setLayout(layout)

    def set_label_text(self, label: QLabel, path: str, prefix: str):
        filename = os.path.basename(path)
        label.setText(f"{prefix} {filename}")
        label.setToolTip(path)

    def select_md_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择 Markdown 文件", "", "Markdown Files (*.md)")
        if file_path:
            self.md_file = file_path
            self.set_label_text(self.label_md, file_path, "📝 已选 Markdown:")

    def select_output_path(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "选择输出位置", "", "Word 文件 (*.docx)")
        if file_path:
            if not file_path.endswith(".docx"):
                file_path += ".docx"
            self.output_file = file_path
            self.set_label_text(self.label_out, file_path, "📁 输出路径:")

    def convert(self):
        if not self.md_file or not self.output_file:
            QMessageBox.warning(self, "缺少信息", "请先选择 Markdown 文件和输出位置")
            return
        try:
            with open(self.md_file, 'r', encoding='utf-8') as f:
                md_text = f.read()
            markdown_to_docx(md_text, self.output_file)
            QMessageBox.information(self, "成功", f"✅ 转换完成！\n输出位置：\n{self.output_file}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"转换失败：\n{e}")


if __name__ == "__main__":
    import warnings
    warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")

    app = QApplication(sys.argv)
    win = MarkdownToWordApp()
    win.show()
    sys.exit(app.exec())
